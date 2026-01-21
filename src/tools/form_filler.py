import os
from pathlib import Path

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("[Warning] python-docx not installed. Word form filling will be limited.")

try:
    from PyPDF2 import PdfReader, PdfWriter
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    print("[Warning] PyPDF2 not installed. PDF form filling will be limited.")

class FormFillerTool:
    def __init__(self):
        self.templates_dir = Path("templates")
        self.output_dir = Path("output")
        
        self.templates_dir.mkdir(exist_ok=True)
        self.output_dir.mkdir(exist_ok=True)

    def fill_word_form(self, template_path: str, data: dict, output_path: str) -> str:
        """
        Fill a Word document template with user data.
        Replaces placeholders like {{name}}, {{student_id}}, etc.
        """
        if not DOCX_AVAILABLE:
            return self._create_text_fallback(template_path, data, output_path)
        
        try:
            doc = Document(template_path)
            
            for paragraph in doc.paragraphs:
                for key, value in data.items():
                    placeholder = f"{{{{{key}}}}}"
                    if placeholder in paragraph.text:
                        paragraph.text = paragraph.text.replace(placeholder, str(value))
            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for key, value in data.items():
                            placeholder = f"{{{{{key}}}}}"
                            if placeholder in cell.text:
                                cell.text = cell.text.replace(placeholder, str(value))
            
            doc.save(output_path)
            return output_path
        
        except Exception as e:
            print(f"[Error] Word form filling failed: {e}")
            return self._create_text_fallback(template_path, data, output_path)

    def fill_pdf_form(self, template_path: str, data: dict, output_path: str) -> str:
        """
        Fill a PDF form with user data.
        This requires PDF form fields to be properly defined in the template.
        """
        if not PDF_AVAILABLE:
            return self._create_text_fallback(template_path, data, output_path)
        
        try:
            reader = PdfReader(template_path)
            writer = PdfWriter()
            
            if '/AcroForm' in reader.trailer['/Root']:
                writer.append_pages_from_reader(reader)
                writer.update_page_form_field_values(
                    writer.pages[0],
                    data
                )
                
                with open(output_path, 'wb') as output_file:
                    writer.write(output_file)
                
                return output_path
            else:
                print("[Warning] PDF does not contain fillable form fields")
                return self._create_text_fallback(template_path, data, output_path)
        
        except Exception as e:
            print(f"[Error] PDF form filling failed: {e}")
            return self._create_text_fallback(template_path, data, output_path)

    def fill_form(self, template_type: str, data: dict, output_path: str = None) -> str:
        """
        Universal form filling interface.
        Detects template type and calls appropriate handler.
        
        Args:
            template_type: Type identifier (e.g., 'scholarship', 'visa', 'enrollment')
            data: Dictionary of field values
            output_path: Optional custom output path
        
        Returns:
            Path to the filled form
        """
        template_extensions = ['.docx', '.pdf', '.hwp', '.doc']
        template_path = None
        
        for ext in template_extensions:
            candidate = self.templates_dir / f"{template_type}{ext}"
            if candidate.exists():
                template_path = str(candidate)
                break
        
        if not template_path:
            print(f"[Warning] Template not found for {template_type}")
            return self._create_text_fallback(None, data, output_path or f"output/{template_type}_filled.txt")
        
        if not output_path:
            ext = Path(template_path).suffix
            output_path = str(self.output_dir / f"{template_type}_filled{ext}")
        
        if template_path.endswith('.docx') or template_path.endswith('.doc'):
            return self.fill_word_form(template_path, data, output_path)
        elif template_path.endswith('.pdf'):
            return self.fill_pdf_form(template_path, data, output_path)
        elif template_path.endswith('.hwp'):
            print("[Warning] HWP format requires hwplib or external conversion")
            return self._create_text_fallback(template_path, data, output_path.replace('.hwp', '.txt'))
        else:
            return self._create_text_fallback(template_path, data, output_path)

    def _create_text_fallback(self, template_path: str, data: dict, output_path: str) -> str:
        """
        Fallback method to create a simple text-based filled form.
        """
        output_path = output_path.replace('.docx', '.txt').replace('.pdf', '.txt').replace('.hwp', '.txt')
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(f"Filled Form\n")
            if template_path:
                f.write(f"Based on template: {template_path}\n")
            f.write("\n" + "="*50 + "\n\n")
            
            for key, value in data.items():
                f.write(f"{key}: {value}\n")
        
        return output_path

    def list_available_templates(self) -> list:
        """List all available form templates"""
        templates = []
        for file in self.templates_dir.iterdir():
            if file.suffix in ['.docx', '.pdf', '.hwp', '.doc']:
                templates.append({
                    'name': file.stem,
                    'format': file.suffix,
                    'path': str(file)
                })
        return templates